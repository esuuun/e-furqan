import docx

try:
    doc = docx.Document('Nabi Muhammad ﷺ.docx')
    with open('docx_output.txt', 'w', encoding='utf-8') as f:
        f.write("--- PARAGRAPHS ---\n")
        for i, p in enumerate(doc.paragraphs[:100]):
            text = p.text.strip()
            if text:
                f.write(f"[{i}] (Style: {p.style.name}) {text}\n")
        
        f.write("\n--- TABLES ---\n")
        for i, table in enumerate(doc.tables[:2]):
            f.write(f"Table {i}:\n")
            for r, row in enumerate(table.rows[:2]):
                for c, cell in enumerate(row.cells[:2]):
                    f.write(f"  Row {r}, Col {c}: {cell.text.strip()[:50]}\n")
except Exception as e:
    with open('docx_output.txt', 'w', encoding='utf-8') as f:
        f.write(f"Error: {e}")
